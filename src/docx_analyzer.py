"""DOCX 文档分析模块 - 用于预览和格式识别"""

from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import FONT_SIZE_MAP


@dataclass
class ParagraphInfo:
    """段落信息"""
    index: int
    text: str
    style_name: str
    font_name: Optional[str]
    font_size: Optional[float]  # 磅值
    font_size_name: Optional[str]  # 字号名称（如小四）
    bold: bool
    italic: bool
    alignment: str
    first_line_indent: Optional[float]
    line_spacing: Optional[float]
    is_heading: bool
    heading_level: int
    format_signature: str  # 格式签名，用于分组


@dataclass  
class FormatGroup:
    """格式分组"""
    signature: str
    font_name: Optional[str]
    font_size: Optional[float]
    font_size_name: Optional[str]
    bold: bool
    italic: bool
    alignment: str
    paragraph_indices: List[int]
    sample_text: str
    suggested_type: str  # 建议的类型：heading1, heading2, body, caption等
    original_type: str = ""  # 原始自动识别的类型


class DocxAnalyzer:
    """DOCX 文档分析器"""
    
    def __init__(self):
        self.paragraphs: List[ParagraphInfo] = []
        self.format_groups: Dict[str, FormatGroup] = {}
        self.document = None
        self.file_path = None
    
    def load_document(self, file_path: str) -> bool:
        """加载文档
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            是否加载成功
        """
        try:
            self.file_path = Path(file_path)
            self.document = Document(file_path)
            self._analyze_paragraphs()
            self._group_by_format()
            return True
        except Exception as e:
            print(f"加载文档失败: {e}")
            return False
    
    def _analyze_paragraphs(self):
        """分析所有段落"""
        self.paragraphs = []
        
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 获取样式信息
            style_name = para.style.name if para.style else "Normal"
            
            # 获取字体信息（优先从runs获取）
            font_name = None
            font_size = None
            bold = False
            italic = False
            
            if para.runs:
                run = para.runs[0]
                if run.font.name:
                    font_name = run.font.name
                if run.font.size:
                    font_size = run.font.size.pt
                bold = run.font.bold or False
                italic = run.font.italic or False
            
            # 如果runs没有，从样式获取
            if para.style and para.style.font:
                if font_name is None and para.style.font.name:
                    font_name = para.style.font.name
                if font_size is None and para.style.font.size:
                    font_size = para.style.font.size.pt
                if not bold and para.style.font.bold:
                    bold = True
            
            # 获取对齐方式
            alignment = "left"
            if para.paragraph_format.alignment:
                align_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: "left",
                    WD_ALIGN_PARAGRAPH.CENTER: "center",
                    WD_ALIGN_PARAGRAPH.RIGHT: "right",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                }
                alignment = align_map.get(para.paragraph_format.alignment, "left")
            
            # 获取首行缩进
            first_line_indent = None
            if para.paragraph_format.first_line_indent:
                first_line_indent = para.paragraph_format.first_line_indent.pt
            
            # 获取行间距
            line_spacing = None
            if para.paragraph_format.line_spacing:
                line_spacing = para.paragraph_format.line_spacing
            
            # 判断是否是标题
            is_heading = style_name.startswith("Heading") or "标题" in style_name
            heading_level = 0
            if is_heading:
                try:
                    heading_level = int(style_name.replace("Heading ", "").replace("标题 ", ""))
                except:
                    heading_level = 1
            
            # 将磅值转换为字号名称
            font_size_name = self._pt_to_size_name(font_size)
            
            # 生成格式签名
            format_signature = self._generate_signature(font_name, font_size, bold, italic, alignment)
            
            para_info = ParagraphInfo(
                index=i,
                text=text[:100] + "..." if len(text) > 100 else text,
                style_name=style_name,
                font_name=font_name,
                font_size=font_size,
                font_size_name=font_size_name,
                bold=bold,
                italic=italic,
                alignment=alignment,
                first_line_indent=first_line_indent,
                line_spacing=line_spacing,
                is_heading=is_heading,
                heading_level=heading_level,
                format_signature=format_signature
            )
            
            self.paragraphs.append(para_info)
    
    def _pt_to_size_name(self, pt: Optional[float]) -> Optional[str]:
        """将磅值转换为字号名称"""
        if pt is None:
            return None
        
        # 查找最接近的字号
        closest_name = None
        min_diff = float('inf')
        
        for name, size in FONT_SIZE_MAP.items():
            diff = abs(size - pt)
            if diff < min_diff:
                min_diff = diff
                closest_name = name
        
        # 如果差距小于0.5磅，返回字号名称
        if min_diff <= 0.5:
            return closest_name
        return f"{pt}磅"
    
    def _generate_signature(self, font_name, font_size, bold, italic, alignment) -> str:
        """生成格式签名"""
        parts = [
            font_name or "default",
            f"{font_size:.1f}" if font_size else "default",
            "B" if bold else "",
            "I" if italic else "",
            alignment
        ]
        return "|".join(str(p) for p in parts)
    
    def _group_by_format(self):
        """按格式分组"""
        self.format_groups = {}
        groups = defaultdict(list)
        
        for para in self.paragraphs:
            groups[para.format_signature].append(para)
        
        # 创建格式组
        for sig, paras in groups.items():
            first_para = paras[0]
            
            # 猜测类型
            suggested_type = self._guess_type(first_para, len(paras))
            
            self.format_groups[sig] = FormatGroup(
                signature=sig,
                font_name=first_para.font_name,
                font_size=first_para.font_size,
                font_size_name=first_para.font_size_name,
                bold=first_para.bold,
                italic=first_para.italic,
                alignment=first_para.alignment,
                paragraph_indices=[p.index for p in paras],
                sample_text=first_para.text[:50],
                suggested_type=suggested_type,
                original_type=suggested_type  # 保存原始类型
            )
    
    def _guess_type(self, para: ParagraphInfo, count: int) -> str:
        """猜测段落类型"""
        # 如果已经是标题样式
        if para.is_heading:
            return f"heading{para.heading_level}"
        
        # 根据字体大小和格式猜测
        if para.font_size:
            if para.font_size >= 15 and para.bold:
                return "heading1"
            elif para.font_size >= 14 and para.bold:
                return "heading2"
            elif para.font_size >= 12 and para.bold:
                return "heading3"
            elif para.font_size <= 9:
                return "caption"
        
        # 居中的小字可能是图表标题
        if para.alignment == "center" and para.font_size and para.font_size <= 10:
            return "caption"
        
        # 默认为正文
        return "body"
    
    def get_paragraphs_by_format(self, signature: str) -> List[ParagraphInfo]:
        """获取指定格式的所有段落"""
        return [p for p in self.paragraphs if p.format_signature == signature]
    
    def get_format_summary(self) -> List[Dict[str, Any]]:
        """获取格式摘要"""
        summary = []
        
        for sig, group in self.format_groups.items():
            summary.append({
                "signature": sig,
                "font_name": group.font_name,
                "font_size": group.font_size,
                "font_size_name": group.font_size_name,
                "bold": group.bold,
                "italic": group.italic,
                "alignment": group.alignment,
                "count": len(group.paragraph_indices),
                "sample": group.sample_text,
                "suggested_type": group.suggested_type
            })
        
        # 按数量排序
        summary.sort(key=lambda x: x["count"], reverse=True)
        
        return summary
    
    def get_preview_html(self) -> str:
        """生成预览HTML"""
        html_parts = ['<html><head><style>']
        html_parts.append('''
            body { font-family: 'Microsoft YaHei', sans-serif; padding: 20px; }
            .paragraph { margin: 10px 0; padding: 8px; border-radius: 4px; cursor: pointer; }
            .paragraph:hover { background-color: #e3f2fd; }
            .paragraph.selected { background-color: #bbdefb; border: 2px solid #1976d2; }
            .format-tag { font-size: 10px; color: #666; background: #f0f0f0; padding: 2px 6px; border-radius: 3px; margin-left: 8px; }
            .heading1 { font-size: 18px; font-weight: bold; }
            .heading2 { font-size: 16px; font-weight: bold; }
            .heading3 { font-size: 14px; font-weight: bold; }
            .body { font-size: 12px; }
            .caption { font-size: 10px; text-align: center; color: #666; }
        ''')
        html_parts.append('</style></head><body>')
        
        for para in self.paragraphs:
            css_class = para.format_signature.replace("|", "_").replace(".", "_")
            type_class = self.format_groups[para.format_signature].suggested_type
            
            style_parts = []
            if para.font_name:
                style_parts.append(f"font-family: '{para.font_name}'")
            if para.font_size:
                style_parts.append(f"font-size: {para.font_size}pt")
            if para.bold:
                style_parts.append("font-weight: bold")
            if para.italic:
                style_parts.append("font-style: italic")
            if para.alignment:
                style_parts.append(f"text-align: {para.alignment}")
            
            style = "; ".join(style_parts)
            
            format_info = f"{para.font_name or '默认'} | {para.font_size_name or '默认'}"
            if para.bold:
                format_info += " | 加粗"
            
            html_parts.append(
                f'<div class="paragraph {type_class}" data-index="{para.index}" '
                f'data-signature="{para.format_signature}" style="{style}">'
                f'{para.text}'
                f'<span class="format-tag">{format_info}</span>'
                f'</div>'
            )
        
        html_parts.append('</body></html>')
        return ''.join(html_parts)
    
    def assign_type_to_format(self, signature: str, element_type: str):
        """将格式分配为特定类型
        
        Args:
            signature: 格式签名
            element_type: 元素类型（heading1, heading2, body, caption等）
        """
        if signature in self.format_groups:
            self.format_groups[signature].suggested_type = element_type
    
    def get_format_mapping(self) -> Dict[str, str]:
        """获取格式映射（签名 -> 类型）"""
        return {sig: group.suggested_type for sig, group in self.format_groups.items()}
    
    def get_all_format_groups(self) -> List[FormatGroup]:
        """获取所有格式组"""
        return list(self.format_groups.values())
