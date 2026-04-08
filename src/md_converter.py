"""Markdown 转 DOCX 转换模块"""

import binascii
import re
import base64
from pathlib import Path
from io import BytesIO
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, UnidentifiedImageError

from .diagnostics import get_logger, log_exception
from .errors import ResourceAccessError
from .resource_policy import ResourcePolicy
from .style_utils import (
    apply_alignment,
    apply_line_spacing,
    apply_run_style,
    apply_runs_style,
    get_font_size_value,
    get_style,
    merge_styles,
)


class MarkdownConverter:
    """Markdown 转换器"""
    
    def __init__(
        self,
        styles: dict = None,
        type_overrides: dict = None,
        resource_policy: ResourcePolicy | None = None,
    ):
        self.styles = merge_styles(styles)
        self.type_overrides = dict(type_overrides or {})
        self.supported_extensions = ['.md', '.markdown']
        self.resource_policy = resource_policy or ResourcePolicy()
        self.base_dir = Path.cwd().resolve()
        self.logger = get_logger("markdown")
    
    def configure(
        self,
        styles: dict = None,
        type_overrides: dict = None,
        resource_policy: ResourcePolicy | None = None,
    ):
        """更新转换器运行时配置。"""
        self.styles = merge_styles(styles)
        self.type_overrides = dict(type_overrides or {})
        if resource_policy is not None:
            self.resource_policy = resource_policy
    
    def convert_to_docx(self, input_path: str, output_path: str = None,
                        progress_callback=None, styles: dict = None,
                        type_overrides: dict = None,
                        resource_policy: ResourcePolicy | None = None) -> str:
        """将Markdown转换为DOCX
        
        Args:
            input_path: Markdown文件路径
            output_path: 输出DOCX文件路径（可选）
            progress_callback: 进度回调函数
            styles: 自定义样式（可选）
            
        Returns:
            输出文件路径
        """
        input_path = Path(input_path)
        if not input_path.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")
        
        if input_path.suffix.lower() not in self.supported_extensions:
            raise ValueError(f"不支持的文件格式: {input_path.suffix}")
        
        if output_path is None:
            output_path = input_path.with_suffix('.docx')
        else:
            output_path = Path(output_path)
        
        self.configure(
            styles=styles,
            type_overrides=type_overrides,
            resource_policy=resource_policy,
        )
        
        # 读取Markdown内容
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 获取Markdown文件所在目录，用于解析相对路径图片
        self.base_dir = input_path.parent.resolve()
        
        if progress_callback:
            progress_callback(10, "解析Markdown内容...")
        
        # 转换为DOCX
        doc = self._md_to_docx(md_content, progress_callback)
        
        if progress_callback:
            progress_callback(90, "保存文档...")
        
        doc.save(str(output_path))
        
        if progress_callback:
            progress_callback(100, "转换完成")
        
        return str(output_path)
    
    def convert_from_string(self, md_content: str, output_path: str,
                            progress_callback=None, styles: dict = None,
                            base_dir: str = None,
                            type_overrides: dict = None,
                            resource_policy: ResourcePolicy | None = None) -> str:
        """从字符串转换Markdown到DOCX
        
        Args:
            md_content: Markdown内容字符串
            output_path: 输出DOCX文件路径
            progress_callback: 进度回调函数
            styles: 自定义样式
            base_dir: 图片基础目录
            
        Returns:
            输出文件路径
        """
        self.configure(
            styles=styles,
            type_overrides=type_overrides,
            resource_policy=resource_policy,
        )
        
        self.base_dir = Path(base_dir).resolve() if base_dir else Path.cwd().resolve()
        
        if progress_callback:
            progress_callback(10, "解析Markdown内容...")
        
        doc = self._md_to_docx(md_content, progress_callback)
        
        if progress_callback:
            progress_callback(90, "保存文档...")
        
        doc.save(str(output_path))
        
        if progress_callback:
            progress_callback(100, "转换完成")
        
        return str(output_path)
    
    def _md_to_docx(self, md_content: str, progress_callback=None) -> Document:
        """将Markdown内容转换为Document对象"""
        doc = Document()
        
        # 设置默认样式
        self._setup_styles(doc)
        
        # 预处理：提取并保护代码块和公式
        code_blocks = []
        formulas = []
        
        # 保护代码块
        def save_code_block(match):
            code_blocks.append(match.group(0))
            return f"<<<CODE_BLOCK_{len(code_blocks) - 1}>>>"
        
        md_content = re.sub(r'```[\s\S]*?```', save_code_block, md_content)
        
        # 保护行内代码
        inline_codes = []
        def save_inline_code(match):
            inline_codes.append(match.group(1))
            return f"<<<INLINE_CODE_{len(inline_codes) - 1}>>>"
        
        md_content = re.sub(r'`([^`]+)`', save_inline_code, md_content)
        
        # 保护公式块
        def save_formula_block(match):
            formulas.append(match.group(0))
            return f"<<<FORMULA_BLOCK_{len(formulas) - 1}>>>"
        
        md_content = re.sub(r'\$\$[\s\S]*?\$\$', save_formula_block, md_content)
        
        # 保护行内公式
        inline_formulas = []
        def save_inline_formula(match):
            inline_formulas.append(match.group(1))
            return f"<<<INLINE_FORMULA_{len(inline_formulas) - 1}>>>"
        
        md_content = re.sub(r'\$([^\$]+)\$', save_inline_formula, md_content)
        
        # 转换为HTML（用于解析复杂结构）
        # 注意：不使用nl2br扩展，避免列表项中的换行产生额外空行
        html = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code', 'toc']
        )
        
        soup = BeautifulSoup(html, 'lxml')
        
        if progress_callback:
            progress_callback(30, "转换文档结构...")
        
        # 处理每个元素
        total_elements = len(soup.find_all(True))
        processed = 0
        
        for element in soup.body.children if soup.body else soup.children:
            self._process_element(doc, element, code_blocks, inline_codes, 
                                 formulas, inline_formulas)
            processed += 1
            if progress_callback and total_elements > 0:
                progress = 30 + int(60 * processed / total_elements)
                progress_callback(min(progress, 90), "转换文档内容...")
        
        return doc
    
    def _get_font_size(self, style_config):
        """获取字体大小（磅值）"""
        return get_font_size_value(style_config)
    
    def _apply_line_spacing(self, paragraph_format, style_config):
        """应用行间距设置"""
        apply_line_spacing(paragraph_format, style_config)
    
    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        styles = doc.styles
        body_style = get_style(self.styles, 'body')
        
        # 设置正文样式
        try:
            normal_style = styles['Normal']
            normal_font = normal_style.font
            
            # 西文字体
            font_en = body_style.get('font_name_en', body_style.get('font_name', 'Times New Roman'))
            normal_font.name = font_en
            
            # 字号
            font_size = self._get_font_size(body_style)
            normal_font.size = Pt(font_size)
            
            # 中文字体
            font_cn = body_style.get('font_name_cn', body_style.get('font_name', '宋体'))
            r_pr = normal_style._element.get_or_add_rPr()
            r_fonts = r_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            r_fonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", font_cn)
        except Exception as error:
            self.logger.warning("应用正文默认样式失败: %s", error)
        
        # 创建各级标题样式
        for i in range(1, 5):
            style_name = f'Heading {i}'
            heading_key = f'heading{i}'
            
            if heading_key in self.styles:
                try:
                    style = styles[style_name]
                    font = style.font
                    heading_style = self.styles[heading_key]
                    
                    # 西文字体
                    font_en = heading_style.get('font_name_en', heading_style.get('font_name', 'Times New Roman'))
                    font.name = font_en
                    
                    # 字号
                    font_size = self._get_font_size(heading_style)
                    font.size = Pt(font_size)
                    
                    font.bold = heading_style.get('bold', True)
                    
                    # 中文字体
                    font_cn = heading_style.get('font_name_cn', heading_style.get('font_name', '宋体'))
                    r_pr = style._element.get_or_add_rPr()
                    r_fonts = r_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
                    if r_fonts is None:
                        r_fonts = OxmlElement("w:rFonts")
                        r_pr.append(r_fonts)
                    r_fonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", font_cn)
                except Exception as error:
                    self.logger.warning("应用标题默认样式失败 | level=%s | error=%s", i, error)

    def _resolve_block_type(self, original_type: str) -> str:
        """根据预览修改结果解析最终块类型。"""
        return self.type_overrides.get(original_type, original_type)

    def _add_text_block(self, doc, text: str, original_type: str, *, list_mode: bool = False):
        """添加文本型块，并允许根据映射重定向样式。"""
        text = text.strip()
        if not text:
            return

        final_type = self._resolve_block_type(original_type)

        if final_type == 'code':
            self._add_code_block(doc, text)
            return

        if final_type == 'formula':
            self._add_formula(doc, text)
            return

        paragraph = doc.add_paragraph(text)

        if final_type.startswith('heading'):
            level = int(final_type[-1]) if final_type[-1].isdigit() else 1
            self._apply_heading_style(paragraph, level)
        elif final_type == 'quote':
            self._apply_quote_style(paragraph)
        elif final_type == 'caption':
            self._apply_caption_style(paragraph)
        elif list_mode:
            self._apply_list_style(paragraph)
        else:
            self._apply_body_style(paragraph)
    
    def _process_element(self, doc, element, code_blocks, inline_codes, 
                        formulas, inline_formulas):
        """处理单个HTML元素"""
        if element.name is None:
            # 纯文本
            text = str(element).strip()
            if text:
                # 恢复特殊内容
                text = self._restore_special_content(
                    text, code_blocks, inline_codes, formulas, inline_formulas
                )
                if text.strip():
                    p = doc.add_paragraph(text)
                    self._apply_body_style(p)
            return
        
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            text = element.get_text()
            text = self._restore_special_content(
                text, code_blocks, inline_codes, formulas, inline_formulas
            )
            self._add_text_block(doc, text, f'heading{min(level, 4)}')
            
        elif element.name == 'p':
            text = element.get_text()
            
            # 检查是否是特殊内容
            if '<<<CODE_BLOCK_' in text:
                match = re.search(r'<<<CODE_BLOCK_(\d+)>>>', text)
                if match:
                    idx = int(match.group(1))
                    self._add_text_block(doc, code_blocks[idx], 'code')
                    return
            
            if '<<<FORMULA_BLOCK_' in text:
                match = re.search(r'<<<FORMULA_BLOCK_(\d+)>>>', text)
                if match:
                    idx = int(match.group(1))
                    self._add_text_block(doc, formulas[idx], 'formula')
                    return
            
            # 检查是否包含图片
            img = element.find('img')
            if img:
                self._add_image(doc, img.get('src', ''), img.get('alt', ''))
                return
            
            text = self._restore_special_content(
                text, code_blocks, inline_codes, formulas, inline_formulas
            )
            self._add_text_block(doc, text, 'body')
                
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                text = li.get_text()
                # 清理多余的空白和换行
                text = ' '.join(text.split())
                text = self._restore_special_content(
                    text, code_blocks, inline_codes, formulas, inline_formulas
                )
                self._add_text_block(doc, text, 'body', list_mode=True)
                
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                text = li.get_text()
                # 清理多余的空白和换行
                text = ' '.join(text.split())
                text = self._restore_special_content(
                    text, code_blocks, inline_codes, formulas, inline_formulas
                )
                self._add_text_block(doc, text, 'body', list_mode=True)
                
        elif element.name == 'blockquote':
            text = element.get_text()
            text = self._restore_special_content(
                text, code_blocks, inline_codes, formulas, inline_formulas
            )
            self._add_text_block(doc, text, 'quote')
            
        elif element.name == 'pre':
            code = element.find('code')
            if code:
                self._add_text_block(doc, code.get_text(), 'code')
            else:
                self._add_text_block(doc, element.get_text(), 'code')
                
        elif element.name == 'table':
            self._add_table(doc, element)
            
        elif element.name == 'img':
            self._add_image(doc, element.get('src', ''), element.get('alt', ''))
            
        elif element.name == 'hr':
            # 添加分隔线
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif element.name in ['div', 'section', 'article']:
            # 递归处理容器元素
            for child in element.children:
                self._process_element(doc, child, code_blocks, inline_codes,
                                     formulas, inline_formulas)
    
    def _restore_special_content(self, text, code_blocks, inline_codes, 
                                 formulas, inline_formulas):
        """恢复特殊内容（代码、公式）"""
        # 恢复代码块
        for i, code in enumerate(code_blocks):
            text = text.replace(f'<<<CODE_BLOCK_{i}>>>', '')
        
        # 恢复行内代码
        for i, code in enumerate(inline_codes):
            text = text.replace(f'<<<INLINE_CODE_{i}>>>', f'「{code}」')
        
        # 恢复公式块
        for i, formula in enumerate(formulas):
            text = text.replace(f'<<<FORMULA_BLOCK_{i}>>>', '')
        
        # 恢复行内公式
        for i, formula in enumerate(inline_formulas):
            text = text.replace(f'<<<INLINE_FORMULA_{i}>>>', f'[公式: {formula}]')
        
        return text
    
    def _apply_body_style(self, paragraph):
        """应用正文样式"""
        style = get_style(self.styles, 'body')
        pf = paragraph.paragraph_format
        
        # 行距
        self._apply_line_spacing(pf, style)
        
        # 段前段后间距
        pf.space_before = Pt(style.get('space_before', 0))
        pf.space_after = Pt(style.get('space_after', 0))
        
        # 首行缩进
        indent = style.get('first_line_indent', 2)
        if indent > 0:
            font_size = self._get_font_size(style)
            pf.first_line_indent = Pt(font_size * indent)
        
        # 对齐方式
        apply_alignment(pf, style.get('alignment', 'left'))
        apply_runs_style(paragraph.runs, style)

    def _apply_list_style(self, paragraph):
        """应用列表样式（无首行缩进，无段前段后间距）"""
        style = get_style(self.styles, 'body')
        pf = paragraph.paragraph_format
        
        # 行距
        self._apply_line_spacing(pf, style)
        
        # 列表项不设置段前段后间距，避免空行
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        
        # 列表项不设置首行缩进
        pf.first_line_indent = Pt(0)
        
        # 对齐方式
        apply_alignment(pf, style.get('alignment', 'left'))
        apply_runs_style(paragraph.runs, style)
    
    def _apply_heading_style(self, heading, level):
        """应用标题样式"""
        style_key = f'heading{min(level, 4)}'
        style = get_style(self.styles, style_key)
        
        pf = heading.paragraph_format
        pf.space_before = Pt(style.get('space_before', 12))
        pf.space_after = Pt(style.get('space_after', 6))
        self._apply_line_spacing(pf, style)
        
        # 对齐方式
        apply_alignment(pf, style.get('alignment', 'left'))
        apply_runs_style(heading.runs, style, bold=style.get('bold', True))
    
    def _apply_quote_style(self, paragraph):
        """应用引用样式"""
        style = get_style(self.styles, 'quote')
        pf = paragraph.paragraph_format
        
        pf.left_indent = Cm(style.get('left_indent', 1))
        pf.space_before = Pt(style.get('space_before', 6))
        pf.space_after = Pt(style.get('space_after', 6))
        self._apply_line_spacing(pf, style)
        apply_alignment(pf, style.get('alignment', 'left'))
        apply_runs_style(paragraph.runs, style, italic=style.get('italic', True))

    def _apply_caption_style(self, paragraph):
        """应用图表标题样式。"""
        style = get_style(self.styles, 'caption')
        pf = paragraph.paragraph_format
        pf.space_before = Pt(style.get('space_before', 6))
        pf.space_after = Pt(style.get('space_after', 6))
        self._apply_line_spacing(pf, style)
        apply_alignment(pf, style.get('alignment', 'center'))
        apply_runs_style(paragraph.runs, style, bold=style.get('bold', False))
    
    def _add_code_block(self, doc, code_text):
        """添加代码块"""
        style = get_style(self.styles, 'code')
        
        # 清理代码文本
        if code_text.startswith('```'):
            lines = code_text.split('\n')
            # 移除首尾的 ```
            lines = lines[1:-1] if lines[-1].strip() == '```' else lines[1:]
            code_text = '\n'.join(lines)
        
        # 创建代码段落
        font_size = self._get_font_size(style) if style.get('font_size') else 10
        font_name = style.get('font_name_en', style.get('font_name', 'Consolas'))
        
        for line in code_text.split('\n'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            apply_run_style(
                run,
                style,
                default_cn='Consolas',
                default_en=font_name,
                default_size=font_size,
            )
            
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            self._apply_line_spacing(pf, style)
            
            # 添加背景色（通过底纹）
            self._add_shading(p, style.get('background', '#f5f5f5'))
    
    def _add_shading(self, paragraph, color):
        """为段落添加底纹"""
        if color.startswith('#'):
            color = color[1:]
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        paragraph._p.get_or_add_pPr().append(shading)
    
    def _add_formula(self, doc, formula_text):
        """添加公式"""
        style = get_style(self.styles, 'formula')
        
        # 清理公式文本
        formula_text = formula_text.strip()
        if formula_text.startswith('$$'):
            formula_text = formula_text[2:]
        if formula_text.endswith('$$'):
            formula_text = formula_text[:-2]
        formula_text = formula_text.strip()
        
        p = doc.add_paragraph()
        run = p.add_run(f'[公式: {formula_text}]')
        run.font.name = 'Cambria Math'
        run.font.size = Pt(12)
        
        pf = p.paragraph_format
        apply_alignment(pf, style.get('alignment', 'center'))
        pf.space_before = Pt(style.get('space_before', 6))
        pf.space_after = Pt(style.get('space_after', 6))
        self._apply_line_spacing(pf, style)
    
    def _add_image(self, doc, src, alt=''):
        """添加图片"""
        style = self.styles.get('image', {})
        
        try:
            # 判断图片来源
            if src.startswith('data:image'):
                # Base64 图片
                image_data = self._decode_base64_image(src)
            else:
                # 本地图片
                image_data = self._resolve_local_image(src)
            
            # 添加图片
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(style.get('space_before', 6))
            pf.space_after = Pt(style.get('space_after', 6))
            
            run = p.add_run()
            
            if isinstance(image_data, str):
                # 文件路径
                run.add_picture(image_data, width=Cm(style.get('max_width', 15)))
            else:
                # BytesIO
                run.add_picture(image_data, width=Cm(style.get('max_width', 15)))
            
            # 添加图片说明（使用caption样式）
            if alt:
                caption_p = doc.add_paragraph(alt)
                self._apply_caption_style(caption_p)
                    
        except ResourceAccessError as error:
            self.logger.warning(
                "图片已跳过 | code=%s | src=%s | base_dir=%s | reason=%s",
                error.code,
                src,
                self.base_dir,
                error.user_message,
            )
            self._add_image_placeholder(doc, alt or src, self._get_image_placeholder_prefix(error))
        except Exception as error:
            log_exception(
                self.logger,
                "图片处理失败",
                error,
                src=src,
                alt=alt,
                base_dir=str(self.base_dir),
            )
            self._add_image_placeholder(doc, alt or src, "图片处理失败")
    
    def _decode_base64_image(self, data_url):
        """解码Base64图片"""
        if "," not in data_url:
            raise ResourceAccessError(
                "嵌入图片数据不完整。",
                code="TODX210",
                hint="请检查 Markdown 中的 data:image 写法是否完整。",
            )

        header, encoded = data_url.split(",", 1)
        mime_type = header[5:].split(";", 1)[0].lower()
        if ";base64" not in header:
            raise ResourceAccessError(
                "仅支持 Base64 编码的嵌入图片。",
                code="TODX211",
                hint="请改用标准的 data:image/...;base64,... 格式。",
            )

        self.resource_policy.validate_embedded_image_mime(mime_type)

        try:
            image_data = base64.b64decode(encoded, validate=True)
        except (ValueError, binascii.Error) as error:
            raise ResourceAccessError(
                "嵌入图片数据损坏，无法解析。",
                code="TODX212",
                hint="请重新生成或替换该嵌入图片。",
            ) from error

        self.resource_policy.validate_resource_size(len(image_data), "嵌入图片")
        self._validate_image_bytes(image_data, label="嵌入图片")
        return BytesIO(image_data)

    def _resolve_local_image(self, src: str) -> str:
        """解析并校验本地图片。"""

        image_path = self.resource_policy.resolve_local_image(self.base_dir, src)
        self._validate_image_file(image_path)
        return str(image_path)

    @staticmethod
    def _get_image_placeholder_prefix(error: ResourceAccessError) -> str:
        """根据错误类型选择占位说明。"""

        if error.code == "TODX207":
            return "图片未找到"
        if error.code in {"TODX202", "TODX204", "TODX205"}:
            return "图片已阻止"
        return "图片未导入"

    def _add_image_placeholder(self, doc, label: str, prefix: str) -> None:
        """添加图片占位提示。"""

        p = doc.add_paragraph(f"[{prefix}: {label}]")
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _validate_image_file(image_path: Path) -> None:
        """验证本地文件确实是受支持图片。"""

        try:
            with Image.open(image_path) as image:
                if image.format not in {"PNG", "JPEG", "GIF", "BMP", "WEBP"}:
                    raise ResourceAccessError(
                        "图片格式不受支持。",
                        code="TODX213",
                        hint="请改用 PNG、JPG、JPEG、GIF、BMP 或 WEBP 图片。",
                    )
                image.verify()
        except (UnidentifiedImageError, OSError) as error:
            raise ResourceAccessError(
                "引用的本地文件不是有效图片。",
                code="TODX214",
                hint="请确认引用的是图片文件，而不是其他类型文件。",
            ) from error

    @staticmethod
    def _validate_image_bytes(image_data: bytes, *, label: str) -> None:
        """验证字节流图片有效性。"""

        try:
            with Image.open(BytesIO(image_data)) as image:
                if image.format not in {"PNG", "JPEG", "GIF", "BMP", "WEBP"}:
                    raise ResourceAccessError(
                        f"{label} 格式不受支持。",
                        code="TODX215",
                        hint="请改用 PNG、JPG、JPEG、GIF、BMP 或 WEBP 图片。",
                    )
                image.verify()
        except (UnidentifiedImageError, OSError) as error:
            raise ResourceAccessError(
                f"{label} 不是有效图片。",
                code="TODX216",
                hint="请重新生成或替换这张图片。",
            ) from error
    
    def _add_table(self, doc, table_element):
        """添加表格"""
        style = get_style(self.styles, 'table')
        
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # 计算列数
        first_row = rows[0]
        cols = len(first_row.find_all(['th', 'td']))
        
        if cols == 0:
            return
        
        # 创建表格
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < cols:
                    table_cell = table.rows[i].cells[j]
                    table_cell.text = cell.get_text().strip()
                    
                    # 设置字体
                    table_font_size = self._get_font_size(style) if style.get('font_size') else 10
                    table_font_cn = style.get('font_name_cn', style.get('font_name', '宋体'))
                    table_font_en = style.get('font_name_en', table_font_cn)
                    
                    for p in table_cell.paragraphs:
                        for run in p.runs:
                            apply_run_style(
                                run,
                                style,
                                default_cn=table_font_cn,
                                default_en=table_font_en,
                                default_size=table_font_size,
                                bold=cell.name == 'th' and style.get('header_bold', True),
                            )
