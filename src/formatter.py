"""智能排版与格式优化模块"""

import os
from pathlib import Path
from typing import Optional, Dict, Any

from .docx_to_md import DocxToMarkdown
from .md_converter import MarkdownConverter
from .config import DEFAULT_STYLES


class SmartFormatter:
    """智能排版与格式优化器"""
    
    def __init__(self, api_key: str = None):
        self.docx_to_md = DocxToMarkdown()
        self.md_converter = MarkdownConverter()
    
    def format_document(self, input_path: str, output_path: str,
                       styles: Dict[str, Any] = None,
                       use_ai: bool = True,
                       progress_callback=None) -> str:
        """智能格式化文档
        
        支持输入格式: .docx, .md
        输出格式: .docx
        
        Args:
            input_path: 输入文件路径
            output_path: 输出DOCX文件路径
            styles: 自定义样式配置
            use_ai: 是否使用AI优化排版
            progress_callback: 进度回调函数
            
        Returns:
            输出文件路径
        """
        input_path = Path(input_path)
        output_path = Path(output_path)
        
        if not input_path.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")
        
        suffix = input_path.suffix.lower()
        
        if progress_callback:
            progress_callback(5, "读取文件...")
        
        # 步骤1: 将文件转换为Markdown（如果是DOCX）
        if suffix in ['.docx', '.doc']:
            if progress_callback:
                progress_callback(10, "将Word转换为Markdown...")
            
            md_content = self.docx_to_md.convert_to_markdown(str(input_path))
            base_dir = input_path.parent
            
        elif suffix in ['.md', '.markdown']:
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            base_dir = input_path.parent
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
        
        if progress_callback:
            progress_callback(30, "文件读取完成")
        
        # 步骤2: 跳过AI优化（已移除AI模块）
        if progress_callback:
            progress_callback(60, "准备生成文档...")
        
        # 步骤3: 转换为DOCX并应用样式
        if progress_callback:
            progress_callback(65, "生成Word文档...")
        
        # 合并样式
        final_styles = {**DEFAULT_STYLES}
        if styles:
            for key, value in styles.items():
                if key in final_styles:
                    final_styles[key] = {**final_styles[key], **value}
                else:
                    final_styles[key] = value
        
        self.md_converter.styles = final_styles
        
        def inner_progress(p, msg):
            if progress_callback:
                # 将内部进度映射到65-95%
                mapped_progress = 65 + int(p * 0.3)
                progress_callback(mapped_progress, msg)
        
        result = self.md_converter.convert_from_string(
            md_content, 
            str(output_path),
            progress_callback=inner_progress,
            base_dir=str(base_dir)
        )
        
        if progress_callback:
            progress_callback(100, "格式化完成")
        
        return result
    

    def apply_selective_format(self, input_path: str, output_path: str,
                               paragraph_mappings: Dict[int, str],
                               styles: Dict[str, Any] = None,
                               progress_callback=None) -> str:
        """选择性格式化：只修改用户指定的段落，其他保持原样
        
        Args:
            input_path: 输入DOCX文件路径
            output_path: 输出DOCX文件路径
            paragraph_mappings: {段落索引: 类型} 的映射，只有这些段落会被修改
            styles: 样式配置
            progress_callback: 进度回调
            
        Returns:
            输出文件路径
        """
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from .config import get_font_size_pt
        
        if progress_callback:
            progress_callback(10, "读取文档...")
        
        # 打开原始文档
        doc = Document(input_path)
        
        # 合并样式配置
        final_styles = {**DEFAULT_STYLES}
        if styles:
            for key, value in styles.items():
                if key in final_styles:
                    final_styles[key] = {**final_styles[key], **value}
                else:
                    final_styles[key] = value
        
        if progress_callback:
            progress_callback(30, "应用格式修改...")
        
        total = len(paragraph_mappings)
        processed = 0
        
        for para_idx, type_id in paragraph_mappings.items():
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                style = final_styles.get(type_id, final_styles.get('body', {}))
                
                self._apply_style_to_paragraph(para, style, type_id, doc)
                
                processed += 1
                if progress_callback and total > 0:
                    prog = 30 + int(60 * processed / total)
                    progress_callback(prog, f"格式化段落 {processed}/{total}")
        
        if progress_callback:
            progress_callback(95, "保存文档...")
        
        doc.save(output_path)
        
        if progress_callback:
            progress_callback(100, "完成")
        
        return output_path

    def _apply_style_to_paragraph(self, para, style: dict, type_id: str, doc=None):
        """将样式应用到单个段落"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from .config import get_font_size_pt
        
        pf = para.paragraph_format
        
        # 段前段后间距（编号段落不设置额外间距以避免空行）
        is_numbered = self._is_numbered_paragraph(para)
        self._doc = doc  # 保存文档引用供后续使用
        if not is_numbered:
            if 'space_before' in style:
                pf.space_before = Pt(style['space_before'])
            if 'space_after' in style:
                pf.space_after = Pt(style['space_after'])
        else:
            # 编号段落：设置较小的间距避免空行
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
        
        # 行距
        spacing_type = style.get('line_spacing_type', '1.5倍行距')
        spacing_value = style.get('line_spacing_value', 1.5)
        if isinstance(spacing_value, str):
            try:
                spacing_value = float(spacing_value)
            except:
                spacing_value = 1.5
        
        if spacing_type == '固定值':
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(float(spacing_value))
        else:
            pf.line_spacing = float(spacing_value)
        
        # 对齐方式
        alignment = style.get('alignment', 'left')
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        pf.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 首行缩进（正文）
        if type_id == 'body' and 'first_line_indent' in style:
            indent_chars = style['first_line_indent']
            font_size = style.get('font_size', 12)
            if isinstance(font_size, str):
                font_size = get_font_size_pt(font_size)
            pf.first_line_indent = Pt(font_size * indent_chars)
        
        # 字体设置
        font_cn = style.get('font_name_cn', style.get('font_name', '宋体'))
        font_en = style.get('font_name_en', style.get('font_name', 'Times New Roman'))
        font_size = style.get('font_size', 12)
        if isinstance(font_size, str):
            font_size = get_font_size_pt(font_size)
        bold = style.get('bold', False)
        
        for run in para.runs:
            run.font.name = font_en
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        
        # 如果是编号段落，还需要修改编号的字体
        if is_numbered:
            self._apply_numbering_font(para, font_cn, font_en, font_size, bold)

    def _is_numbered_paragraph(self, para) -> bool:
        """检查段落是否为编号段落"""
        from docx.oxml.ns import qn
        pPr = para._element.pPr
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                return True
        return False

    def _apply_numbering_font(self, para, font_cn: str, font_en: str, font_size: float, bold: bool):
        """应用字体到编号 - 通过修改文档的编号定义"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        pPr = para._element.pPr
        if pPr is None:
            return
        
        # 获取编号属性
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return
        
        # 获取 numId 和 ilvl
        numId_elem = numPr.find(qn('w:numId'))
        ilvl_elem = numPr.find(qn('w:ilvl'))
        
        if numId_elem is None:
            return
        
        numId = numId_elem.get(qn('w:val'))
        ilvl = ilvl_elem.get(qn('w:val')) if ilvl_elem is not None else '0'
        
        # 使用保存的文档引用来修改编号定义
        doc = getattr(self, '_doc', None)
        if doc is not None:
            try:
                if hasattr(doc, 'part') and hasattr(doc.part, 'numbering_part') and doc.part.numbering_part:
                    numbering = doc.part.numbering_part.numbering_definitions._numbering
                    
                    # 查找对应的 num 元素
                    for num in numbering.findall(qn('w:num')):
                        if num.get(qn('w:numId')) == numId:
                            abstractNumId = num.find(qn('w:abstractNumId'))
                            if abstractNumId is not None:
                                absNumId = abstractNumId.get(qn('w:val'))
                                # 查找 abstractNum
                                for absNum in numbering.findall(qn('w:abstractNum')):
                                    if absNum.get(qn('w:abstractNumId')) == absNumId:
                                        # 查找对应级别的 lvl
                                        for lvl in absNum.findall(qn('w:lvl')):
                                            if lvl.get(qn('w:ilvl')) == ilvl:
                                                self._update_lvl_font(lvl, font_cn, font_en, font_size, bold)
                                                break
                                        break
                            break
            except Exception:
                pass
        
        # 同时在段落级别设置 rPr 作为备用
        rPr = pPr.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            pPr.insert(0, rPr)
        
        self._update_rPr_font(rPr, font_cn, font_en, font_size, bold)
    
    def _update_lvl_font(self, lvl, font_cn: str, font_en: str, font_size: float, bold: bool):
        """更新编号级别的字体设置"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # 查找或创建 rPr
        rPr = lvl.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            lvl.append(rPr)
        
        self._update_rPr_font(rPr, font_cn, font_en, font_size, bold)
    
    def _update_rPr_font(self, rPr, font_cn: str, font_en: str, font_size: float, bold: bool):
        """更新 rPr 元素的字体设置（包括清除颜色）"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # 清除现有字体设置和颜色
        for child in list(rPr):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ['rFonts', 'sz', 'szCs', 'b', 'bCs', 'color']:
                rPr.remove(child)
        
        # 设置字体
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_en)
        rFonts.set(qn('w:hAnsi'), font_en)
        rFonts.set(qn('w:eastAsia'), font_cn)
        rFonts.set(qn('w:cs'), font_en)
        rPr.append(rFonts)
        
        # 设置字号
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)
        
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(szCs)
        
        # 设置粗体
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)
        
        # 设置颜色为黑色（清除彩色编号）
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')
        rPr.append(color)


class StylePreset:
    """样式预设（使用新格式）"""
    
    # 学术论文样式
    ACADEMIC = {
        "heading1": {
            "font_name_cn": "黑体",
            "font_name_en": "Times New Roman",
            "font_size": "三号",
            "bold": True,
            "space_before": 24,
            "space_after": 12,
            "line_spacing_type": "固定值",
            "line_spacing_value": 20,
            "alignment": "center",
        },
        "heading2": {
            "font_name_cn": "黑体",
            "font_name_en": "Times New Roman",
            "font_size": "四号",
            "bold": True,
            "space_before": 12,
            "space_after": 6,
            "line_spacing_type": "固定值",
            "line_spacing_value": 20,
            "alignment": "left",
        },
        "heading3": {
            "font_name_cn": "黑体",
            "font_name_en": "Times New Roman",
            "font_size": "小四",
            "bold": True,
            "space_before": 6,
            "space_after": 3,
            "line_spacing_type": "固定值",
            "line_spacing_value": 20,
            "alignment": "left",
        },
        "body": {
            "font_name_cn": "宋体",
            "font_name_en": "Times New Roman",
            "font_size": "小四",
            "line_spacing_type": "固定值",
            "line_spacing_value": 20,
            "first_line_indent": 2,
            "alignment": "justify",
        },
        "caption": {
            "font_name_cn": "宋体",
            "font_name_en": "Times New Roman",
            "font_size": "五号",
            "line_spacing_type": "单倍行距",
            "line_spacing_value": 1.0,
            "alignment": "center",
        },
    }
    
    # 公文样式
    OFFICIAL = {
        "heading1": {
            "font_name_cn": "方正小标宋简体",
            "font_name_en": "Times New Roman",
            "font_size": "二号",
            "bold": False,
            "space_before": 0,
            "space_after": 0,
            "line_spacing_type": "固定值",
            "line_spacing_value": 28,
            "alignment": "center",
        },
        "heading2": {
            "font_name_cn": "黑体",
            "font_name_en": "Times New Roman",
            "font_size": "三号",
            "bold": False,
            "space_before": 12,
            "space_after": 6,
            "line_spacing_type": "固定值",
            "line_spacing_value": 28,
            "alignment": "left",
        },
        "heading3": {
            "font_name_cn": "楷体",
            "font_name_en": "Times New Roman",
            "font_size": "三号",
            "bold": False,
            "space_before": 6,
            "space_after": 3,
            "line_spacing_type": "固定值",
            "line_spacing_value": 28,
            "alignment": "left",
        },
        "body": {
            "font_name_cn": "仿宋",
            "font_name_en": "Times New Roman",
            "font_size": "三号",
            "line_spacing_type": "固定值",
            "line_spacing_value": 28,
            "first_line_indent": 2,
            "alignment": "justify",
        },
        "caption": {
            "font_name_cn": "楷体",
            "font_name_en": "Times New Roman",
            "font_size": "小四",
            "line_spacing_type": "固定值",
            "line_spacing_value": 28,
            "alignment": "center",
        },
    }
    
    @classmethod
    def get_preset(cls, name: str) -> dict:
        """获取预设样式"""
        presets = {
            "academic": cls.ACADEMIC,
            "official": cls.OFFICIAL,
        }
        return presets.get(name.lower(), DEFAULT_STYLES)
    
    @classmethod
    def list_presets(cls) -> list:
        """列出所有预设"""
        return ["academic", "official"]
