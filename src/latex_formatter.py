"""LaTeX 格式化模块 - 将 LaTeX 转换为 DOCX"""

import re
from pathlib import Path
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from .latex_analyzer import LatexAnalyzer, LatexParagraphInfo
from .config import DEFAULT_STYLES, FONT_SIZE_MAP


class LatexToDocxConverter:
    """LaTeX 到 DOCX 转换器 - 完全使用传入的样式配置"""
    
    def __init__(self, analyzer: LatexAnalyzer, styles: Dict[str, Any] = None):
        self.analyzer = analyzer
        self.styles = styles or {}
        self.doc = Document()
        # 标题编号计数器：[一级, 二级, 三级, 四级, 五级]
        self.heading_counters = [0, 0, 0, 0, 0]
        # 表格和代码计数器
        self.table_counter = 0
        self.code_counter = 0
    
    def convert(self, paragraph_mappings: Dict[int, str] = None,
                progress_callback=None) -> Document:
        """将 LaTeX 转换为 DOCX"""
        if progress_callback:
            progress_callback(10, "准备转换...")
        
        paragraph_mappings = paragraph_mappings or {}
        total = len(self.analyzer.paragraphs)
        
        for i, para in enumerate(self.analyzer.paragraphs):
            # 确定段落类型：优先使用用户修改的，否则使用原始类型
            if para.index in paragraph_mappings:
                element_type = paragraph_mappings[para.index]
            else:
                element_type = para.original_type
            
            # 根据原始类型选择不同的处理方法
            original_type = para.original_type
            
            if original_type == 'table':
                # 表格：解析并创建 DOCX 表格
                self._add_table(para.raw_text)
            elif original_type == 'code':
                # 代码：保持格式
                self._add_code_block(para.raw_text)
            elif original_type == 'formula':
                # 块级公式
                self._add_formula(para.raw_text)
            else:
                # 普通文本：清理后输出，处理行内公式
                full_text = self._clean_latex_for_docx(para.raw_text)
                self._add_paragraph_with_style(full_text, element_type)
            
            if progress_callback and i % 10 == 0:
                progress = 10 + int(80 * i / total)
                progress_callback(progress, f"转换中... {i}/{total}")
        
        if progress_callback:
            progress_callback(95, "完成转换")
        
        return self.doc
    
    def _add_paragraph_with_style(self, text: str, element_type: str):
        """根据类型添加段落并应用左侧面板的样式，处理行内公式"""
        p = self.doc.add_paragraph()
        
        # 如果是标题，添加编号
        if element_type.startswith('heading'):
            level = int(element_type[-1]) if element_type[-1].isdigit() else 1
            numbered_text = self._get_heading_number(level) + text
            text = numbered_text
        
        # 获取对应类型的样式
        if element_type.startswith('heading'):
            style = self.styles.get(element_type, self.styles.get('body', {}))
        else:
            style = self.styles.get('body', {})
        
        # 字体设置
        font_cn = style.get('font_name_cn') or style.get('font_cn', '宋体')
        font_en = style.get('font_name_en') or style.get('font_en', 'Times New Roman')
        font_size = style.get('font_size', '小四')
        size_pt = self._get_font_size_pt(font_size)
        is_bold = style.get('bold', False)
        
        # 处理行内公式 $...$：分割文本，交替处理普通文本和公式
        parts = re.split(r'(\$[^$]+\$)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('$') and part.endswith('$'):
                # 行内公式：用斜体表示
                formula_text = part[1:-1]  # 去掉 $ 符号
                # 清理公式中的转义字符
                formula_text = self._unescape_latex(formula_text)
                run = p.add_run(formula_text)
                run.font.name = 'Cambria Math'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
                run.font.italic = True
                run.font.size = Pt(size_pt)
            else:
                # 普通文本 - 处理转义字符
                clean_text = self._unescape_latex(part)
                run = p.add_run(clean_text)
                run.font.name = font_en
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
                run.font.bold = is_bold
                run.font.size = Pt(size_pt)
        
        # 首行缩进（正文类型才有）
        if not element_type.startswith('heading'):
            indent = style.get('first_line_indent', 2)
            if indent:
                p.paragraph_format.first_line_indent = Cm(indent * 0.35)
        
        # 行距
        line_type = style.get('line_spacing_type', '倍数行距')
        line_value = style.get('line_spacing_value', 1.5)
        if line_type == '固定值':
            p.paragraph_format.line_spacing = Pt(line_value)
        else:
            p.paragraph_format.line_spacing = line_value
        
        # 对齐方式
        align = style.get('alignment', 'left')
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'justify':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _get_font_size_pt(self, size_name) -> float:
        """将字号名称转换为磅值"""
        if isinstance(size_name, (int, float)):
            return float(size_name)
        return FONT_SIZE_MAP.get(size_name, 12)
    
    def _get_heading_number(self, level: int) -> str:
        """生成中文风格的标题编号
        一级：一、二、三、...
        二级：1、2、3、...
        三级：1.1、1.2、...
        四级：1.1.1、1.1.2、...
        五级及以下：①②③...
        """
        if level < 1 or level > 5:
            return ""
        
        # 更新计数器
        idx = level - 1
        self.heading_counters[idx] += 1
        
        # 重置下级计数器
        for i in range(idx + 1, len(self.heading_counters)):
            self.heading_counters[i] = 0
        
        # 生成编号字符串
        if level == 1:
            # 一级：一、二、三、...
            return f"{self._to_chinese_number(self.heading_counters[0])}、"
        elif level == 2:
            # 二级：1、2、3、...
            return f"{self.heading_counters[1]}. "
        elif level == 3:
            # 三级：1.1、1.2、...
            return f"{self.heading_counters[1]}.{self.heading_counters[2]} "
        elif level == 4:
            # 四级：1.1.1、1.1.2、...
            return f"{self.heading_counters[1]}.{self.heading_counters[2]}.{self.heading_counters[3]} "
        else:
            # 五级：①②③...
            return f"{self._to_circled_number(self.heading_counters[4])}"
    
    def _to_chinese_number(self, num: int) -> str:
        """将数字转换为中文数字"""
        chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                       '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']
        if 0 <= num <= 20:
            return chinese_nums[num]
        elif num < 100:
            tens = num // 10
            ones = num % 10
            if ones == 0:
                return f"{chinese_nums[tens]}十"
            else:
                return f"{chinese_nums[tens]}十{chinese_nums[ones]}"
        return str(num)
    
    def _to_circled_number(self, num: int) -> str:
        """将数字转换为带圈数字 ①②③..."""
        circled = ['⓪', '①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩',
                   '⑪', '⑫', '⑬', '⑭', '⑮', '⑯', '⑰', '⑱', '⑲', '⑳']
        if 0 <= num <= 20:
            return circled[num]
        return f"({num})"
    
    def _add_caption(self, text: str):
        """添加图表/代码标题，使用 caption 样式，上下各有一个空行"""
        # 获取 caption 样式
        style = self.styles.get('caption', {})
        font_cn = style.get('font_name_cn') or style.get('font_cn', '宋体')
        font_en = style.get('font_name_en') or style.get('font_en', 'Times New Roman')
        font_size = style.get('font_size', '小五')
        size_pt = self._get_font_size_pt(font_size)
        
        # 上方空行
        p_before = self.doc.add_paragraph()
        run_before = p_before.add_run()
        run_before.font.size = Pt(size_pt)
        
        # 标题行
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = font_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        run.font.size = Pt(size_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 下方空行
        p_after = self.doc.add_paragraph()
        run_after = p_after.add_run()
        run_after.font.size = Pt(size_pt)
    
    def _add_table(self, raw_text: str):
        """解析 LaTeX 表格并创建 DOCX 表格"""
        # 提取 caption（稍后在表格下方添加）
        caption_match = re.search(r'\\caption\{([^}]*)\}', raw_text)
        caption_text = None
        if caption_match:
            caption_text = self._clean_cell(caption_match.group(1))
        
        # 提取 tabular 内容
        tabular_match = re.search(r'\\begin\{tabular\}\{[^}]*\}(.*?)\\end\{tabular\}', raw_text, re.DOTALL)
        if not tabular_match:
            # 如果没有 tabular，尝试其他格式
            return
        
        tabular_content = tabular_match.group(1)
        
        # 先移除 \toprule, \midrule, \bottomrule, \hline, \cline 等规则
        tabular_content = re.sub(r'\\(?:toprule|midrule|bottomrule|hline|cline\{[^}]*\})\s*', '', tabular_content)
        
        # 解析表格行（按 \\ 分割，但要注意 \\ 在 LaTeX 中是换行）
        rows_raw = re.split(r'\s*\\\\\s*', tabular_content)
        rows = []
        for row in rows_raw:
            row = row.strip()
            if not row:
                continue
            # 按 & 分割单元格
            cells = [self._clean_cell(c.strip()) for c in row.split('&')]
            if any(c for c in cells):  # 至少有一个非空单元格
                rows.append(cells)
        
        if not rows:
            return
        
        # 创建 DOCX 表格
        num_cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            row_obj = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row_obj.cells[j]
                    cell.text = cell_text
                    # 设置单元格字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.size = Pt(10.5)
        
        # 在表格下方添加 caption（类似 "表1 xxx" 格式），使用 caption 样式
        if caption_text:
            self.table_counter += 1
            self._add_caption(f"表{self.table_counter}  {caption_text}")
        
        # 表格后添加空行
        self.doc.add_paragraph()
    
    def _clean_cell(self, text: str) -> str:
        """清理表格单元格内容"""
        # 移除 \textbf 等格式命令，保留内容
        text = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\textit\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\texttt\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\emph\{([^}]*)\}', r'\1', text)
        # 移除其他命令
        text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\[a-zA-Z]+(?![_])', '', text)  # 不匹配 \_
        text = re.sub(r'[{}]', '', text)
        # 处理转义字符
        text = self._unescape_latex(text)
        return text.strip()
    
    def _add_code_block(self, raw_text: str):
        """添加代码块，保持格式"""
        # 提取代码内容（去掉 \begin{lstlisting}[...] 和 \end{lstlisting}）
        code_match = re.search(r'\\begin\{(?:lstlisting|verbatim|minted)\}(?:\[[^\]]*\])?\s*(.*?)\s*\\end\{(?:lstlisting|verbatim|minted)\}', 
                               raw_text, re.DOTALL)
        if code_match:
            code_content = code_match.group(1)
        else:
            # 直接使用原始内容
            code_content = raw_text
        
        # 提取 caption（稍后在代码块下方添加）
        caption_match = re.search(r'caption=([^,\]]+)', raw_text)
        caption_text = None
        if caption_match:
            caption_text = self._unescape_latex(caption_match.group(1).strip())
        
        # 获取代码样式
        code_style = self.styles.get('code', {})
        code_font = code_style.get('font_name', 'Consolas')
        code_size = self._get_font_size_pt(code_style.get('font_size', '小五'))
        
        # 按行添加代码，保持缩进
        for line in code_content.split('\n'):
            if line.strip() or line:  # 保留空行
                p = self.doc.add_paragraph()
                run = p.add_run(line.rstrip())
                # 使用等宽字体
                run.font.name = code_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                run.font.size = Pt(code_size)
                # 不缩进
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.left_indent = Cm(0.5)
        
        # 在代码块下方添加 caption（类似 "代码1 xxx" 格式），使用 caption 样式
        if caption_text:
            self.code_counter += 1
            self._add_caption(f"代码{self.code_counter}  {caption_text}")
    
    def _add_formula(self, raw_text: str):
        """添加块级公式 - 使用 Word 公式对象"""
        from docx.oxml.ns import nsmap
        from docx.oxml import OxmlElement
        
        # 提取公式内容
        formula_match = re.search(r'\\begin\{(?:equation|align|gather|multline)\*?\}(.*?)\\end\{(?:equation|align|gather|multline)\*?\}', 
                                  raw_text, re.DOTALL)
        if formula_match:
            formula_content = formula_match.group(1).strip()
        else:
            formula_content = raw_text
        
        # 清理公式内容
        formula_content = self._clean_formula(formula_content)
        
        # 创建段落
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 创建 OMML 公式对象
        self._insert_omml_formula(p, formula_content)
    
    def _clean_formula(self, text: str) -> str:
        """清理公式内容，转换为 Word 可识别的格式"""
        # 移除 LaTeX 特殊命令，保留数学符号
        text = re.sub(r'\\times', '×', text)
        text = re.sub(r'\\div', '÷', text)
        text = re.sub(r'\\pm', '±', text)
        text = re.sub(r'\\leq', '≤', text)
        text = re.sub(r'\\geq', '≥', text)
        text = re.sub(r'\\neq', '≠', text)
        text = re.sub(r'\\approx', '≈', text)
        text = re.sub(r'\\infty', '∞', text)
        text = re.sub(r'\\alpha', 'α', text)
        text = re.sub(r'\\beta', 'β', text)
        text = re.sub(r'\\gamma', 'γ', text)
        text = re.sub(r'\\delta', 'δ', text)
        text = re.sub(r'\\sum', '∑', text)
        text = re.sub(r'\\prod', '∏', text)
        text = re.sub(r'\\int', '∫', text)
        text = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', text)
        text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text)
        text = re.sub(r'\\log', 'log', text)
        text = re.sub(r'\\ln', 'ln', text)
        text = re.sub(r'\\sin', 'sin', text)
        text = re.sub(r'\\cos', 'cos', text)
        text = re.sub(r'\\tan', 'tan', text)
        # 移除剩余的反斜杠命令
        text = re.sub(r'\\[a-zA-Z]+', '', text)
        text = re.sub(r'[{}]', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def _insert_omml_formula(self, paragraph, formula_text: str):
        """插入 OMML 公式到段落"""
        from lxml import etree
        
        # OMML 命名空间
        OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        
        # 创建 oMathPara (公式段落)
        omath_para = etree.Element('{%s}oMathPara' % OMML_NS)
        
        # 创建 oMath (公式)
        omath = etree.SubElement(omath_para, '{%s}oMath' % OMML_NS)
        
        # 创建 r (run)
        r = etree.SubElement(omath, '{%s}r' % OMML_NS)
        
        # 创建 t (text)
        t = etree.SubElement(r, '{%s}t' % OMML_NS)
        t.text = formula_text
        
        # 将公式添加到段落
        paragraph._p.append(omath_para)
    
    def _clean_latex_for_docx(self, text: str) -> str:
        """清理 LaTeX 命令，提取纯文本用于 DOCX 输出"""
        # 移除环境标记
        text = re.sub(r'\\begin\{\w+\*?\}(\[.*?\])?', '', text)
        text = re.sub(r'\\end\{\w+\*?\}', '', text)
        
        # 提取标题内容
        text = re.sub(r'\\(?:sub)*(?:section|chapter|paragraph)\*?\{([^}]*)\}', r'\1', text)
        
        # 提取网址 - 保留 URL 内容
        text = re.sub(r'\\url\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\href\{([^}]*)\}\{([^}]*)\}', r'\2 (\1)', text)  # 显示文字 (网址)
        
        # 提取格式命令内容
        text = re.sub(r'\\(?:textbf|textit|emph|underline|textrm|textsf|texttt)\{([^}]*)\}', r'\1', text)
        
        # 提取 caption 内容
        text = re.sub(r'\\caption\{([^}]*)\}', r'\1', text)
        
        # 移除 \item 标记
        text = re.sub(r'\\item\s*(\[[^\]]*\])?\s*', '', text)
        
        # 移除引用和标签
        text = re.sub(r'\\label\{[^}]*\}', '', text)
        text = re.sub(r'\\ref\{[^}]*\}', '[ref]', text)
        text = re.sub(r'\\cite\{[^}]*\}', '[cite]', text)
        
        # 移除其他常见命令
        text = re.sub(r'\\(?:centering|raggedright|raggedleft|noindent|par)\b', '', text)
        text = re.sub(r'\\(?:vspace|hspace)\*?\{[^}]*\}', '', text)
        text = re.sub(r'\\(?:small|large|Large|LARGE|huge|Huge|tiny|footnotesize|scriptsize|normalsize)\b', '', text)
        
        # 移除剩余的简单命令（但保留 $...$）
        text = re.sub(r'\\[a-zA-Z]+\*?(\{[^}]*\})?', '', text)
        
        # 清理多余的大括号
        text = re.sub(r'[{}]', '', text)
        
        # 处理 LaTeX 转义字符 - 必须在清理命令之后
        text = self._unescape_latex(text)
        
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def _unescape_latex(self, text: str) -> str:
        """处理 LaTeX 转义字符"""
        # 注意：这里使用字面字符串，不是正则表达式
        # LaTeX 转义字符在文本中的形式
        escapes = [
            ('\\_', '_'),   # 下划线
            ('\\%', '%'),   # 百分号
            ('\\&', '&'),   # 与号
            ('\\#', '#'),   # 井号
            ('\\~', '~'),   # 波浪号
            ('\\^', '^'),   # 脱字符
            ('\\{', '{'),   # 左大括号
            ('\\}', '}'),   # 右大括号
            ('\\$', '$'),   # 美元符号
        ]
        for latex, char in escapes:
            text = text.replace(latex, char)
        return text


def convert_latex_to_docx(input_file: str, output_file: str,
                          paragraph_mappings: Dict[int, str] = None,
                          styles: Dict[str, Any] = None,
                          progress_callback=None) -> str:
    """将 LaTeX 文件转换为 DOCX
    
    Args:
        input_file: 输入的 .tex 文件路径
        output_file: 输出的 .docx 文件路径
        paragraph_mappings: 段落类型映射
        styles: 样式配置
        progress_callback: 进度回调
        
    Returns:
        输出文件路径
    """
    # 加载分析器
    analyzer = LatexAnalyzer()
    if not analyzer.load_document(input_file):
        raise ValueError(f"无法加载 LaTeX 文件: {input_file}")
    
    # 转换
    converter = LatexToDocxConverter(analyzer, styles)
    doc = converter.convert(paragraph_mappings, progress_callback)
    
    # 保存
    doc.save(output_file)
    
    if progress_callback:
        progress_callback(100, "保存完成")
    
    return output_file


class LatexFormatter:
    """LaTeX 格式化器"""
    
    # 类型到 LaTeX 标题命令的映射
    TYPE_TO_HEADING = {
        'heading1': r'\section',
        'heading2': r'\subsection',
        'heading3': r'\subsubsection',
        'heading4': r'\paragraph',
    }
    
    # 字号映射
    FONT_SIZE_MAP = {
        '初号': r'\Huge',
        '小初': r'\huge',
        '一号': r'\LARGE',
        '小一': r'\Large',
        '二号': r'\large',
        '小二': r'\large',
        '三号': r'\normalsize',
        '小三': r'\normalsize',
        '四号': r'\normalsize',
        '小四': r'\small',
        '五号': r'\footnotesize',
        '小五': r'\scriptsize',
        '六号': r'\tiny',
    }
    
    def __init__(self, analyzer: LatexAnalyzer):
        self.analyzer = analyzer
        self.lines = analyzer.lines.copy()
    
    def apply_format(self, paragraph_mappings: Dict[int, str],
                     styles: Dict[str, Any] = None,
                     progress_callback=None) -> str:
        """应用格式修改
        
        Args:
            paragraph_mappings: {段落索引: 目标类型} 的映射
            styles: 样式配置
            progress_callback: 进度回调
            
        Returns:
            修改后的 LaTeX 内容
        """
        if progress_callback:
            progress_callback(10, "分析修改内容...")
        
        # 按行号倒序处理，避免行号偏移
        modifications = []
        for para_idx, target_type in paragraph_mappings.items():
            para = self.analyzer.get_paragraph_by_index(para_idx)
            if para:
                modifications.append((para, target_type, styles))
        
        # 按起始行号倒序排序
        modifications.sort(key=lambda x: x[0].line_start, reverse=True)
        
        total = len(modifications)
        processed = 0
        
        for para, target_type, styles in modifications:
            self._modify_paragraph(para, target_type, styles)
            processed += 1
            if progress_callback and total > 0:
                prog = 10 + int(80 * processed / total)
                progress_callback(prog, f"修改段落 {processed}/{total}")
        
        if progress_callback:
            progress_callback(95, "生成输出...")
        
        return '\n'.join(self.lines)
    
    def _modify_paragraph(self, para: LatexParagraphInfo, 
                          target_type: str, styles: Dict[str, Any]):
        """修改单个段落"""
        style = styles.get(target_type, {}) if styles else {}
        
        # 获取原始行
        start = para.line_start
        end = para.line_end
        original_lines = self.lines[start:end + 1]
        original_text = '\n'.join(original_lines)
        
        # 根据目标类型生成新内容
        if target_type.startswith('heading'):
            new_content = self._format_as_heading(original_text, para, target_type, style)
        elif target_type == 'body':
            new_content = self._format_as_body(original_text, para, style)
        elif target_type == 'caption':
            new_content = self._format_as_caption(original_text, para, style)
        elif target_type == 'code':
            new_content = self._format_as_code(original_text, para, style)
        elif target_type == 'quote':
            new_content = self._format_as_quote(original_text, para, style)
        else:
            new_content = original_text
        
        # 替换行
        new_lines = new_content.split('\n')
        self.lines[start:end + 1] = new_lines
    
    def _format_as_heading(self, text: str, para: LatexParagraphInfo,
                           target_type: str, style: dict) -> str:
        """格式化为标题"""
        # 提取纯文本内容
        content = self._extract_content(text, para.original_type)
        
        # 获取对应的标题命令
        heading_cmd = self.TYPE_TO_HEADING.get(target_type, r'\section')
        
        # 构建新的标题行
        result = f"{heading_cmd}{{{content}}}"
        
        # 如果需要加粗但标题默认不加粗，可以包装
        if style.get('bold', True):
            # LaTeX 标题默认加粗，不需要额外处理
            pass
        
        return result
    
    def _format_as_body(self, text: str, para: LatexParagraphInfo,
                        style: dict) -> str:
        """格式化为正文"""
        content = self._extract_content(text, para.original_type)
        
        result_parts = []
        
        # 首行缩进（LaTeX 默认有缩进，如果不要可以用 \noindent）
        indent = style.get('first_line_indent', 2)
        if indent == 0:
            result_parts.append(r'\noindent ')
        
        # 加粗
        if style.get('bold', False):
            content = f"\\textbf{{{content}}}"
        
        result_parts.append(content)
        
        return ''.join(result_parts)
    
    def _format_as_caption(self, text: str, para: LatexParagraphInfo,
                           style: dict) -> str:
        """格式化为图表标题"""
        # 如果已经是 figure/table 环境，保持不变
        if r'\begin{figure}' in text or r'\begin{table}' in text:
            return text
        
        content = self._extract_content(text, para.original_type)
        
        # 简单地添加居中
        return f"\\begin{{center}}\n{content}\n\\end{{center}}"
    
    def _format_as_code(self, text: str, para: LatexParagraphInfo,
                        style: dict) -> str:
        """格式化为代码"""
        content = self._extract_content(text, para.original_type)
        
        # 如果已经是代码环境，保持不变
        if r'\begin{verbatim}' in text or r'\begin{lstlisting}' in text:
            return text
        
        return f"\\begin{{verbatim}}\n{content}\n\\end{{verbatim}}"
    
    def _format_as_quote(self, text: str, para: LatexParagraphInfo,
                         style: dict) -> str:
        """格式化为引用"""
        content = self._extract_content(text, para.original_type)
        
        # 如果已经是引用环境，保持不变
        if r'\begin{quote}' in text or r'\begin{quotation}' in text:
            return text
        
        return f"\\begin{{quotation}}\n{content}\n\\end{{quotation}}"
    
    def _extract_content(self, text: str, original_type: str) -> str:
        """从原始文本中提取内容"""
        # 如果是标题，提取花括号内的内容
        if original_type.startswith('heading'):
            match = re.search(r'\\(?:sub)*(?:section|chapter|paragraph)\*?\{([^}]*)\}', text)
            if match:
                return match.group(1)
        
        # 如果是环境，提取环境内容
        env_match = re.match(r'\\begin\{(\w+)\}(.*?)\\end\{\1\}', text, re.DOTALL)
        if env_match:
            return env_match.group(2).strip()
        
        # 否则返回清理后的文本
        # 移除标题命令
        text = re.sub(r'\\(?:sub)*(?:section|chapter|paragraph)\*?\{([^}]*)\}', r'\1', text)
        return text.strip()
    
    def save(self, output_path: str, content: str = None):
        """保存到文件"""
        if content is None:
            content = '\n'.join(self.lines)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return output_path


def format_latex_file(input_path: str, output_path: str,
                      paragraph_mappings: Dict[int, str],
                      styles: Dict[str, Any] = None,
                      progress_callback=None) -> str:
    """格式化 LaTeX 文件的便捷函数
    
    Args:
        input_path: 输入 .tex 文件路径
        output_path: 输出 .tex 文件路径
        paragraph_mappings: {段落索引: 目标类型} 的映射
        styles: 样式配置
        progress_callback: 进度回调
        
    Returns:
        输出文件路径
    """
    analyzer = LatexAnalyzer()
    if not analyzer.load_document(input_path):
        raise ValueError(f"无法加载 LaTeX 文件: {input_path}")
    
    formatter = LatexFormatter(analyzer)
    content = formatter.apply_format(paragraph_mappings, styles, progress_callback)
    formatter.save(output_path, content)
    
    if progress_callback:
        progress_callback(100, "完成")
    
    return output_path
