"""安全边界与请求校验测试。"""

from __future__ import annotations

import base64
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from PIL import Image

from src.conversion_request import build_conversion_request
from src.errors import ValidationError
from src.formatter import SmartFormatter
from src.md_converter import MarkdownConverter
from src.resource_policy import ResourcePolicy
from src.user_settings import UserSettingsStore


def _create_png(path: Path, size: tuple[int, int] = (8, 8), color: tuple[int, int, int] = (255, 0, 0)) -> None:
    image = Image.new("RGB", size, color)
    image.save(path, format="PNG")


class SecurityBoundaryTests(unittest.TestCase):
    def test_markdown_remote_image_is_blocked(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "sample.md"
            output_path = tmp_path / "sample.docx"
            input_path.write_text("![远程图](https://example.com/demo.png)", encoding="utf-8")

            SmartFormatter().format_document(
                str(input_path),
                str(output_path),
                use_ai=False,
            )

            doc = Document(output_path)
            self.assertIn("[图片已阻止: 远程图]", [paragraph.text for paragraph in doc.paragraphs])

    def test_markdown_local_image_cannot_escape_base_dir(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            docs_dir = tmp_path / "docs"
            docs_dir.mkdir()
            outside_image = tmp_path / "outside.png"
            _create_png(outside_image)

            output_path = tmp_path / "blocked.docx"
            MarkdownConverter().convert_from_string(
                "![越界图](../outside.png)",
                str(output_path),
                base_dir=str(docs_dir),
            )

            doc = Document(output_path)
            self.assertIn("[图片已阻止: 越界图]", [paragraph.text for paragraph in doc.paragraphs])

    def test_markdown_local_image_allows_safe_relative_file(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            image_path = tmp_path / "demo.png"
            output_path = tmp_path / "safe.docx"
            _create_png(image_path)

            MarkdownConverter().convert_from_string(
                "![本地图](demo.png)",
                str(output_path),
                base_dir=str(tmp_path),
            )

            doc = Document(output_path)
            self.assertEqual(len(doc.inline_shapes), 1)
            self.assertIn("本地图", [paragraph.text for paragraph in doc.paragraphs])

    def test_embedded_image_respects_size_limit(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            image_path = tmp_path / "embedded.png"
            output_path = tmp_path / "embedded.docx"
            _create_png(image_path, size=(32, 32))
            encoded = base64.b64encode(image_path.read_bytes()).decode("ascii")
            markdown_text = f"![嵌入图](data:image/png;base64,{encoded})"

            converter = MarkdownConverter(resource_policy=ResourcePolicy(max_image_bytes=32))
            converter.convert_from_string(
                markdown_text,
                str(output_path),
                base_dir=str(tmp_path),
            )

            doc = Document(output_path)
            self.assertIn("[图片未导入: 嵌入图]", [paragraph.text for paragraph in doc.paragraphs])

    def test_build_conversion_request_uses_input_relative_output_dir(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "sample.md"
            input_path.write_text("# demo", encoding="utf-8")

            request = build_conversion_request(
                str(input_path),
                "exports",
                {"body": {"font_name_cn": "宋体"}},
            )

            self.assertEqual(request.output_dir, (tmp_path / "exports").resolve())
            self.assertTrue(request.output_dir.exists())
            self.assertEqual(request.output_path.name, "sample_formatted.docx")

    def test_build_conversion_request_rejects_output_path_that_is_file(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "sample.md"
            output_file = tmp_path / "not-a-dir.txt"
            input_path.write_text("# demo", encoding="utf-8")
            output_file.write_text("x", encoding="utf-8")

            with self.assertRaises(ValidationError):
                build_conversion_request(
                    str(input_path),
                    str(output_file),
                    {"body": {"font_name_cn": "宋体"}},
                )

    def test_user_settings_round_trip_last_output_dir(self):
        with TemporaryDirectory() as tmp:
            settings_path = Path(tmp) / "settings.json"
            store = UserSettingsStore(settings_path)

            store.update_last_output_dir(r"E:\Output")
            settings = store.load()

            self.assertEqual(settings.last_output_dir, r"E:\Output")


if __name__ == "__main__":
    unittest.main()
