"""核心转换链路的回归测试。"""

from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.formatter import SmartFormatter
from src.latex_formatter import convert_latex_to_docx
from src.style_utils import merge_styles


class ConversionStabilityTests(unittest.TestCase):
    def test_merge_styles_preserves_nested_defaults(self):
        merged = merge_styles({"body": {"font_name_cn": "仿宋"}})

        self.assertEqual(merged["body"]["font_name_cn"], "仿宋")
        self.assertIn("first_line_indent", merged["body"])
        self.assertIn("line_spacing_type", merged["body"])

    def test_markdown_type_override_reaches_converter(self):
        styles = {
            "heading1": {
                "font_name_cn": "黑体",
                "font_name_en": "Times New Roman",
                "font_size": "三号",
                "bold": True,
                "alignment": "center",
            },
            "body": {
                "font_name_cn": "宋体",
                "font_name_en": "Times New Roman",
                "font_size": "小四",
                "bold": False,
                "alignment": "justify",
                "first_line_indent": 2,
            },
        }

        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "sample.md"
            output_path = tmp_path / "sample.docx"
            input_path.write_text("# 标题\n\n正文段落", encoding="utf-8")

            SmartFormatter().format_document(
                str(input_path),
                str(output_path),
                styles=styles,
                type_overrides={"heading1": "body"},
                use_ai=False,
            )

            doc = Document(output_path)
            first_paragraph = doc.paragraphs[0]

            self.assertEqual(first_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertFalse(first_paragraph.runs[0].font.bold)
            self.assertIsNotNone(first_paragraph.paragraph_format.first_line_indent)

    def test_latex_paragraph_mapping_uses_target_style(self):
        styles = {
            "body": {
                "font_name_cn": "宋体",
                "font_name_en": "Times New Roman",
                "font_size": "小四",
                "alignment": "left",
                "first_line_indent": 2,
            },
            "caption": {
                "font_name_cn": "黑体",
                "font_name_en": "Times New Roman",
                "font_size": "小五",
                "alignment": "center",
                "space_before": 6,
                "space_after": 6,
            },
        }

        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_path = tmp_path / "sample.tex"
            output_path = tmp_path / "sample.docx"
            input_path.write_text(
                "\\documentclass{article}\n"
                "\\begin{document}\n"
                "\\section{章节}\n"
                "这是正文。\n"
                "\\end{document}\n",
                encoding="utf-8",
            )

            convert_latex_to_docx(
                str(input_path),
                str(output_path),
                paragraph_mappings={1: "caption"},
                styles=styles,
            )

            doc = Document(output_path)
            body_paragraph = doc.paragraphs[1]

            self.assertEqual(body_paragraph.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertIsNone(body_paragraph.paragraph_format.first_line_indent)


if __name__ == "__main__":
    unittest.main()
